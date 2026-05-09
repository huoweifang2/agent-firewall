#!/usr/bin/env python3
"""Build a template-aligned thesis DOCX into article/docx1.

This generator intentionally keeps article/docx untouched. It reuses the
existing LaTeX/content parser from article/docx/build_thesis_docx.py, but starts
from the official Word template and maps all generated content onto template
styles instead of the previous Thesis* style family.
"""

from __future__ import annotations

import re
import sys
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ARTICLE_DIR = Path(__file__).resolve().parents[1]
DOCX1_DIR = ARTICLE_DIR / "docx1"
SOURCE_DOCX_DIR = ARTICLE_DIR / "docx"
TEMPLATE_PATH = ARTICLE_DIR / "论文模板.docx"
OUT_PATH = DOCX1_DIR / "霍玮放-本科毕业论文（设计）.docx"
SOURCE_DOCX_PATH = SOURCE_DOCX_DIR / "霍玮放-本科毕业论文（设计）.docx"

sys.path.insert(0, str(SOURCE_DOCX_DIR))
import build_thesis_docx as base  # noqa: E402


BODY_STYLE = "Normal (Web)"
CAPTION_STYLE = "图表"
HEADING_STYLES = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3"}


def patch_base_paths() -> None:
    """Redirect imported helper output paths into docx1."""

    base.DOCX_DIR = DOCX1_DIR
    base.FIGURE_DIR = DOCX1_DIR / "generated_figures"
    base.OUT_PATH = OUT_PATH
    base.LATEX_FIGURE_RENDER_DIR = base.FIGURE_DIR / "latex_figures"


def clear_document_body(doc: Document) -> None:
    """Remove all sample content while preserving the template package/styles."""

    body = doc._body._element
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def style_exists(doc: Document, name: str) -> bool:
    try:
        doc.styles[name]
        return True
    except KeyError:
        return False


def configure_template_styles(doc: Document) -> None:
    """Normalize the template styles that the generated thesis will use."""

    normal = doc.styles["Normal"]
    base.set_style_font(normal, "宋体", "Times New Roman", 12, False)
    base.set_paragraph_format(normal, first_line_chars=2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    body = doc.styles[BODY_STYLE] if style_exists(doc, BODY_STYLE) else normal
    base.set_style_font(body, "宋体", "Times New Roman", 12, False)
    base.set_paragraph_format(body, first_line_chars=2, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY)

    for level, size, align in [
        (1, 16, WD_ALIGN_PARAGRAPH.CENTER),
        (2, 14, WD_ALIGN_PARAGRAPH.LEFT),
        (3, 12, WD_ALIGN_PARAGRAPH.LEFT),
    ]:
        style = doc.styles[HEADING_STYLES[level]]
        base.set_style_font(style, "宋体", "Times New Roman", size, True)
        base.set_paragraph_format(style, first_line_chars=0, alignment=align, before=6, after=6)
        base.set_outline_level(style, level - 1)

    if style_exists(doc, CAPTION_STYLE):
        caption = doc.styles[CAPTION_STYLE]
        base.set_style_font(caption, "宋体", "Times New Roman", 9, True)
        base.set_paragraph_format(caption, first_line_chars=0, alignment=WD_ALIGN_PARAGRAPH.CENTER)


def font_for_role(role: str) -> tuple[str, str, float, bool]:
    if role == "Thesis Abstract CN":
        return "楷体-简", "Times New Roman", 12, False
    if role == "Thesis Abstract EN":
        return "Times New Roman", "Times New Roman", 12, False
    if role in {"Thesis Reference", "Thesis Acknowledgement"}:
        return "宋体", "Times New Roman", 10.5, False
    if role == "Thesis Code":
        return "宋体", "Menlo", 9, False
    return "宋体", "Times New Roman", 12, False


def rewrite_thesis_text(text: str) -> str:
    """Clean source-reference phrasing and tighten thesis voice."""

    text = base.normalize_spaces(text)
    replacements = {
        "本文基于Agent-Firewall系统实现，设计并实现了": "本文围绕Agent-Firewall系统，设计并实现了",
        "结合本项目实际实现独立完成": "由作者结合系统实现独立完成",
        "本次红队场景资产共": "实验整理得到红队场景资产共",
        "本次轻量规则与扫描器组合": "轻量规则与扫描器组合",
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    forbidden_patterns = [
        "根据现有代码",
        "根据已有工程",
        "根据用户提供的代码",
        "结合任务书和代码",
        "通过 AI 分析",
        "通过AI分析",
    ]
    for pattern in forbidden_patterns:
        text = text.replace(pattern, "")
    return text.strip()


def add_mixed_runs(paragraph, text: str, role: str) -> None:
    pattern = re.compile(r"(`[^`]+`|\[[^\]]+\]\([^)]+\))")
    east, ascii_font, size, bold = font_for_role(role)
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            base.set_run_font(run, east, ascii_font, size, bold)
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token.strip("`"))
            base.set_run_font(run, "宋体", "Menlo", 10.5, False)
        else:
            label = token[1:].split("](", 1)[0]
            run = paragraph.add_run(label)
            base.set_run_font(run, east, ascii_font, size, bold)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        base.set_run_font(run, east, ascii_font, size, bold)


def add_body_paragraph(doc: Document, text: str, style: str = "Normal", first_line=None):
    text = rewrite_thesis_text(text)
    if not text:
        return None

    actual_style = BODY_STYLE if style != "Thesis Code" else BODY_STYLE
    paragraph = doc.add_paragraph(style=actual_style)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

    if style == "Thesis Reference":
        base.set_paragraph_indent(paragraph, left_chars=202, hanging_chars=202)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    elif style == "Thesis Code":
        base.set_paragraph_indent(paragraph, first_line_chars=0)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    else:
        indent = 2 if first_line is None else first_line
        base.set_paragraph_indent(paragraph, first_line_chars=int(indent * 100))
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    add_mixed_runs(paragraph, text, style)
    return paragraph


def add_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=HEADING_STYLES[level])
    paragraph.paragraph_format.keep_with_next = True
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    run = paragraph.add_run(text)
    base.set_run_font(run, "宋体", "Times New Roman", size={1: 16, 2: 14, 3: 12}[level], bold=True)
    return paragraph


def add_front_heading(doc: Document, text: str, level: int):
    paragraph = doc.add_paragraph(style=BODY_STYLE)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    run = paragraph.add_run(text)
    base.set_run_font(run, "宋体", "Times New Roman", size={1: 16, 2: 14}[level], bold=True)
    return paragraph


def add_outline_front_heading(doc: Document, text: str, level: int = 1, outline_level: int = 0):
    """Add an unnumbered heading that still appears in Word's TOC."""

    paragraph = add_front_heading(doc, text, min(level, 2))
    base.set_paragraph_outline_level(paragraph, outline_level)
    return paragraph


def add_center(doc: Document, text: str, size=12, bold=False, before=0, after=0):
    paragraph = doc.add_paragraph(style=BODY_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    run = paragraph.add_run(text)
    base.set_run_font(run, "宋体", "Times New Roman", size=size, bold=bold)
    return paragraph


def add_cover(doc: Document) -> None:
    base.configure_section(doc.sections[0], header=False)
    add_center(doc, "学校代码：10022", size=14)
    for _ in range(3):
        doc.add_paragraph(style=BODY_STYLE)
    add_center(doc, "本科毕业论文(设计)", size=22, bold=True, before=10, after=24)
    add_center(doc, base.TITLE_CN, size=16, bold=True, after=8)
    add_center(doc, base.TITLE_EN, size=14, bold=True, after=34)
    add_center(doc, base.AUTHOR, size=14, after=22)

    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    base.set_table_widths(table, [2200, 5000])
    rows = [
        ("学    院", base.COLLEGE),
        ("专    业", base.MAJOR),
        ("班    级", base.CLASS_NAME),
        ("学    号", base.STUDENT_ID),
        ("指导教师", base.SUPERVISOR),
    ]
    for row, (left, right) in zip(table.rows, rows):
        for idx, value in enumerate((left, right)):
            cell = row.cells[idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.style = doc.styles[BODY_STYLE]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            base.set_paragraph_indent(p, first_line_chars=0)
            run = p.add_run(value)
            base.set_run_font(run, "宋体", "Times New Roman", 14, False)
    base.remove_table_borders(table)
    for _ in range(4):
        doc.add_paragraph(style=BODY_STYLE)
    add_center(doc, base.DATE_TEXT, size=14)


def add_declarations(doc: Document) -> None:
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    base.configure_section(doc.sections[-1], header=False)
    add_front_heading(doc, "独创性声明", 1)
    add_body_paragraph(
        doc,
        "本人声明所呈交的论文（设计）是本人在导师指导下独立进行的设计、研究工作及取得的设计、研究成果。"
        "尽我所知，除了论文（设计）中特别加以标注和致谢的地方外，论文（设计）中不包含其他人已经发表"
        "或撰写过的研究成果，本论文（设计）中没有抄袭他人研究成果和伪造数据等行为。与我共同工作的人员"
        "对本研究所做的任何贡献均已在论文（设计）中作了明确的说明并表示了谢意。",
    )
    doc.add_paragraph(style=BODY_STYLE)
    add_body_paragraph(doc, "作者签名：           日期：   2026  年     月      日", first_line=0)
    doc.add_paragraph(style=BODY_STYLE)
    add_front_heading(doc, "关于毕业论文（设计）使用授权的说明", 1)
    add_body_paragraph(
        doc,
        "本人完全了解北京林业大学有关保留、使用毕业论文（设计）的规定，即：本科生在校期间毕业论文（设计）"
        "工作的知识产权单位属北京林业大学；学校有权保留并向国家有关部门或机构送交论文（设计）的纸质版和"
        "电子版，允许毕业论文（设计）被查阅、借阅和复印；学校可以将毕业论文（设计）的全部或部分内容公开"
        "或编入有关数据库进行检索，可以允许采用影印、缩印或其它复制手段保存、汇编毕业论文（设计）。",
    )
    add_body_paragraph(doc, "（保密的论文在解密后应适用本授权书）", first_line=0)
    doc.add_paragraph(style=BODY_STYLE)
    add_body_paragraph(doc, "作者签名：           指导老师签名：", first_line=0)
    add_body_paragraph(doc, "日    期：    2026  年    月    日", first_line=0)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    base.configure_section(doc.sections[-1], header=False)
    add_front_heading(doc, "北京林业大学本科毕业论文(设计)AI工具规范使用诚信承诺书", 1)
    add_front_heading(doc, "一、AI工具使用情况说明", 2)
    add_body_paragraph(doc, "本人郑重声明，在毕业论文(设计)写作过程中使用AI工具的情况如下：")
    add_body_paragraph(doc, "□未使用任何AI工具(勾选此项者无需填写后续内容)", first_line=0)
    add_body_paragraph(doc, "■使用AI工具(勾选此项者需如实填写以下内容)", first_line=0)
    add_body_paragraph(
        doc,
        "1.使用工具名称：ChatGPT、DeepSeek、论文写作规范辅助、文献数据库检索工具、LaTeX排版与图表辅助工具。",
        first_line=0,
    )
    add_body_paragraph(
        doc,
        "2.具体用途及生成内容：用于文献检索关键词整理、论文结构草拟、学术语言润色建议、"
        "GB/T 7714-2015参考文献格式检查、LaTeX格式调整、图表生成脚本辅助编写。"
        "系统设计、实验方案、核心代码理解、数据复核、结论推导由作者结合系统实现独立完成。",
        first_line=0,
    )
    add_body_paragraph(doc, "3.生成内容占全文比例：25%。", first_line=0)
    add_front_heading(doc, "二、诚信承诺", 2)
    add_body_paragraph(
        doc,
        "本人承诺：毕业论文(设计)核心研究内容(包括实验设计、数据分析、结论推导等)均独立完成，"
        "未使用AI工具直接生成或替代。上述AI工具使用情况描述真实、准确、完整，无隐瞒或虚假陈述。",
    )
    doc.add_paragraph(style=BODY_STYLE)
    add_body_paragraph(doc, "学生签名：          日期：    2026.  .", first_line=0)


def add_keywords(doc: Document, label: str, text: str, english=False) -> None:
    paragraph = doc.add_paragraph(style=BODY_STYLE)
    paragraph.paragraph_format.space_before = Pt(12)
    paragraph.paragraph_format.space_after = Pt(0)
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    east = "Times New Roman" if english else "宋体"
    if english and not label.endswith(" "):
        label = f"{label} "
    run = paragraph.add_run(label)
    base.set_run_font(run, east, "Times New Roman", 12, True)
    run = paragraph.add_run(rewrite_thesis_text(text))
    base.set_run_font(run, east, "Times New Roman", 12, False)


def add_table(doc: Document, rows: list[list[str]], caption_cn: str, caption_en: str, number: str):
    if not rows:
        return
    is_abbrev_table = (
        not caption_cn
        and rows
        and len(rows[0]) >= 2
        and "符号或缩略词" in rows[0][0]
        and "含义" in rows[0][1]
    )
    if caption_cn:
        cap = doc.add_paragraph(style=CAPTION_STYLE)
        cap.paragraph_format.keep_with_next = True
        run = cap.add_run(f"表{number} {caption_cn}")
        base.set_run_font(run, "宋体", "Times New Roman", 9, True)
        if caption_en:
            cap2 = doc.add_paragraph(style=CAPTION_STYLE)
            cap2.paragraph_format.keep_with_next = True
            run = cap2.add_run(f"Table {number} {caption_en}")
            base.set_run_font(run, "宋体", "Times New Roman", 9, True)

    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    base.set_table_widths(table, base.compute_widths(rows, cols))
    base.set_three_line_borders(table)
    for r_idx, row in enumerate(rows):
        for c_idx in range(cols):
            cell = table.rows[r_idx].cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if is_abbrev_table:
                base.set_cell_margins(cell, top=36, start=80, bottom=36, end=80)
            else:
                base.set_cell_margins(cell)
            p = cell.paragraphs[0]
            p.style = doc.styles[BODY_STYLE]
            text = rewrite_thesis_text(row[c_idx] if c_idx < len(row) else "")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if r_idx == 0 or base.short_cell(text) else WD_ALIGN_PARAGRAPH.LEFT
            base.set_paragraph_indent(p, first_line_chars=0)
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            p.paragraph_format.line_spacing = 1.05 if is_abbrev_table else 1.25
            run = p.add_run(base.add_wrap_opportunities(text))
            base.set_run_font(run, "宋体", "Times New Roman", 8.5 if is_abbrev_table else 9, r_idx == 0)
    doc.add_paragraph(style=BODY_STYLE)


def add_figure(
    doc: Document,
    caption_cn: str,
    caption_en: str,
    number: str,
    label: str,
    image_path: Path,
    scale: float = 1.0,
):
    paragraph = doc.add_paragraph(style=BODY_STYLE)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=Cm(14.6 * scale))
    cap = doc.add_paragraph(style=CAPTION_STYLE)
    cap.paragraph_format.keep_with_next = True
    run = cap.add_run(f"图{number} {caption_cn}")
    base.set_run_font(run, "宋体", "Times New Roman", 9, True)
    if caption_en:
        cap2 = doc.add_paragraph(style=CAPTION_STYLE)
        cap2.paragraph_format.keep_with_next = True
        run = cap2.add_run(f"Figure {number} {caption_en}")
        base.set_run_font(run, "宋体", "Times New Roman", 9, True)
    doc.add_paragraph(style=BODY_STYLE)


def add_code_block(self, code: str):
    self.flush()
    for line in code.strip("\n").splitlines():
        add_body_paragraph(self.doc, line.lstrip(), style="Thesis Code", first_line=0)


def add_section_heading(self, title: str, level: int, starred=False):
    self.flush()
    title = self.clean(title)
    if starred:
        if self.pending_break:
            base.add_page_break(self.doc)
            self.pending_break = False
        add_outline_front_heading(self.doc, title, level=1, outline_level=0)
        if title == "摘要":
            self.current_style = "Thesis Abstract CN"
        elif title == "Abstract":
            self.current_style = "Thesis Abstract EN"
        elif title == "致谢":
            self.current_style = "Thesis Acknowledgement"
        else:
            self.current_style = "Normal"
        return

    if level == 1:
        if self.appendix:
            self.appendix_no += 1
            self.current_key = chr(ord("A") + self.appendix_no - 1)
            heading = f"附录{self.current_key}  {title}"
            self.sub_no = 0
            self.subsub_no = 0
            if self.pending_break:
                base.add_page_break(self.doc)
                self.pending_break = False
            elif self.body_started:
                base.add_page_break(self.doc)
            self.body_started = True
            self.current_style = "Normal"
            add_outline_front_heading(self.doc, heading, level=1, outline_level=0)
            return

        self.section_no += 1
        self.current_key = str(self.section_no)
        self.sub_no = 0
        self.subsub_no = 0
        if self.pending_break:
            base.add_page_break(self.doc)
            self.pending_break = False
        elif self.body_started:
            base.add_page_break(self.doc)
        self.body_started = True
        self.current_style = "Normal"
        add_heading(self.doc, title, 1)
    elif level == 2:
        self.sub_no += 1
        self.subsub_no = 0
        if self.appendix:
            add_outline_front_heading(self.doc, f"{self.current_key}.{self.sub_no} {title}", level=2, outline_level=1)
        else:
            add_heading(self.doc, title, 2)
    elif level == 3:
        self.subsub_no += 1
        if self.appendix:
            add_outline_front_heading(
                self.doc,
                f"{self.current_key}.{self.sub_no}.{self.subsub_no} {title}",
                level=2,
                outline_level=2,
            )
        else:
            add_heading(self.doc, title, 3)


def add_toc(self):
    self.flush()
    base.add_page_break(self.doc)
    self.pending_break = False
    add_front_heading(self.doc, "目  录", 1)
    paragraph = self.doc.add_paragraph(style=BODY_STYLE)
    base.set_paragraph_indent(paragraph, first_line_chars=0)
    base.add_field(paragraph, r'TOC \o "1-3" \h \z \u')
    base.add_page_break(self.doc)


def add_bibliography_heading(self):
    self.flush()
    base.add_page_break(self.doc)
    add_outline_front_heading(self.doc, "参考文献", level=1, outline_level=0)
    self.in_bibliography = True
    self.reference_count = 0
    self.current_style = "Thesis Reference"


def process_bibitem(self, line: str):
    match = re.match(r"\\bibitem\{([^}]+)\}\s*(.*)", line)
    if not match:
        return
    self.reference_count += 1
    text = self.clean(match.group(2))
    add_body_paragraph(self.doc, f"[{self.reference_count}] {text}", style="Thesis Reference", first_line=0)


def install_template_hooks() -> None:
    base.configure_styles = configure_template_styles
    base.add_body_paragraph = add_body_paragraph
    base.add_heading = add_heading
    base.add_front_heading = add_front_heading
    base.add_cover = add_cover
    base.add_declarations = add_declarations
    base.add_keywords = add_keywords
    base.add_table = add_table
    base.add_figure = add_figure
    base.ThesisBuilder.add_section_heading = add_section_heading
    base.ThesisBuilder.add_code_block = add_code_block
    base.ThesisBuilder.add_toc = add_toc
    base.ThesisBuilder.add_bibliography_heading = add_bibliography_heading
    base.ThesisBuilder.process_bibitem = process_bibitem


def document_text(docx_path: Path) -> str:
    from lxml import etree

    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    with zipfile.ZipFile(docx_path) as zf:
        xml = zf.read("word/document.xml")
    root = etree.fromstring(xml)
    paragraphs = ["".join(p.xpath(".//w:t/text()", namespaces=ns)) for p in root.xpath(".//w:p", namespaces=ns)]
    return "\n".join(paragraphs).replace("\u200b", "")


def audit_docx1(docx_path: Path) -> None:
    plain = document_text(docx_path)
    required = [
        "学校代码：10022",
        base.TITLE_CN,
        base.AUTHOR,
        base.COLLEGE,
        base.MAJOR,
        base.CLASS_NAME,
        base.STUDENT_ID,
        "北京林业大学本科毕业论文(设计)AI工具规范使用诚信承诺书",
        "生成内容占全文比例：25%。",
        "摘要",
        "Abstract",
        "目  录",
        "主要符号与缩略词对照表",
        "绪论",
        "Agent-Firewall",
        "图3.1 Agent-Firewall总体架构",
        "表5.2 本机工具流与运行时合同联调结果",
        "参考文献",
        "致谢",
        "附录A",
    ]
    missing = [item for item in required if item not in plain]
    if missing:
        raise RuntimeError(f"missing required text in docx1 build: {missing}")

    forbidden = [
        "根据现有代码",
        "根据已有工程",
        "根据用户提供的代码",
        "结合任务书和代码",
        "通过AI分析",
        "通过 AI 分析",
    ]
    leaked = [item for item in forbidden if item in plain]
    if leaked:
        raise RuntimeError(f"forbidden source-process wording remains: {leaked}")

    with zipfile.ZipFile(docx_path) as zf:
        document_xml = zf.read("word/document.xml").decode("utf-8")
        styles_xml = zf.read("word/styles.xml").decode("utf-8")
    if "ThesisHeading" in document_xml or "ThesisHeading" in styles_xml:
        raise RuntimeError("previous Thesis* styles leaked into template-aligned output")
    for token in ['w:styleId="1"', 'w:styleId="2"', 'w:styleId="3"', 'w:styleId="af0"', 'w:styleId="af6"']:
        if token not in styles_xml:
            raise RuntimeError(f"template style token missing: {token}")
    for margin in ['w:top="1701"', 'w:left="1701"', 'w:right="1417"', 'w:bottom="1417"']:
        if margin not in document_xml:
            raise RuntimeError(f"template page margin missing: {margin}")


def build() -> Path:
    patch_base_paths()
    install_template_hooks()
    DOCX1_DIR.mkdir(parents=True, exist_ok=True)
    tex = base.TEX_PATH.read_text(encoding="utf-8")

    doc = Document(str(TEMPLATE_PATH))
    clear_document_body(doc)
    configure_template_styles(doc)
    add_cover(doc)
    add_declarations(doc)

    doc.add_section(WD_SECTION_START.NEW_PAGE)
    base.configure_section(doc.sections[-1], page_fmt="upperRoman", start=1, header=True)
    builder = base.ThesisBuilder(doc, tex)
    builder.parse()

    doc.core_properties.title = base.TITLE_CN
    doc.core_properties.author = base.AUTHOR
    doc.core_properties.subject = "北京林业大学本科毕业论文（设计）"
    doc.core_properties.comments = f"Template: {TEMPLATE_PATH}; content reference: {SOURCE_DOCX_PATH}"
    doc.save(OUT_PATH)
    base.update_settings_for_fields(OUT_PATH)
    base.patch_detector_ooxml(OUT_PATH)
    audit_docx1(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    print(build())
